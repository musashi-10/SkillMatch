import os
import re
from io import BytesIO

CANONICAL_SKILLS = [
    "python", "sql", "javascript", "typescript", "java", "c++", "react", "node",
    "django", "flask", "fastapi", "pandas", "numpy", "scikit-learn", "tensorflow",
    "pytorch", "docker", "kubernetes", "aws", "azure", "gcp", "power bi", "tableau",
    "excel", "machine learning", "deep learning", "nlp", "postgresql", "redis",
]

ROLE_KEYWORDS = {
    "Data Scientist": {"machine learning", "deep learning", "pandas", "numpy", "python"},
    "Data Analyst": {"sql", "excel", "power bi", "tableau", "reporting"},
    "Backend Developer": {"python", "django", "flask", "fastapi", "api", "postgresql"},
    "Frontend Developer": {"react", "javascript", "typescript", "html", "css"},
    "DevOps Engineer": {"docker", "kubernetes", "aws", "terraform", "ci/cd"},
    "Software Engineer": {"python", "java", "c++", "algorithms", "data structures"},
}


def extract_resume_text(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise ValueError("resume file_path does not exist")
    ext = os.path.splitext(file_path)[1].lower()
    if ext in {".txt", ".md"}:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    if ext == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            raise ValueError("PDF parsing requires 'pypdf' package")
        reader = PdfReader(file_path)
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if ext == ".docx":
        try:
            import docx  # type: ignore
        except Exception:
            raise ValueError("DOCX parsing requires 'python-docx' package")
        doc = docx.Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValueError("Unsupported resume format. Use txt, md, pdf, or docx")


def extract_resume_bytes(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename or "")[1].lower()
    if ext in {".txt", ".md"}:
        return data.decode("utf-8", errors="ignore")
    if ext == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            raise ValueError("PDF parsing requires 'pypdf' package")
        reader = PdfReader(BytesIO(data))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if ext == ".docx":
        try:
            import docx  # type: ignore
        except Exception:
            raise ValueError("DOCX parsing requires 'python-docx' package")
        doc = docx.Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    raise ValueError("Unsupported resume format. Use txt, md, pdf, or docx")


def extract_skills_and_roles(text: str) -> dict:
    text_l = (text or "").lower()
    skills = []
    for s in CANONICAL_SKILLS:
        if re.search(rf"\b{re.escape(s)}\b", text_l):
            skills.append(s)
    roles = []
    skill_set = set(skills)
    for role, kws in ROLE_KEYWORDS.items():
        if len(skill_set & kws) >= 2:
            roles.append(role)
    return {"skills": skills[:30], "roles": roles[:6]}
